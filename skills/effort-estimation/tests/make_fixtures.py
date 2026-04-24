#!/usr/bin/env python3
"""Generate the two binary test fixtures (input.docx and input.xlsx)."""

from pathlib import Path
from docx import Document
from openpyxl import Workbook

HERE = Path(__file__).parent


def make_docx():
    doc = Document()
    doc.add_heading("需求：企业客户关系管理（CRM）子模块", level=1)
    doc.add_paragraph(
        "为销售团队提供客户线索、跟进记录、商机管理能力。本次需求为独立子模块，挂在现有中台上，"
        "不涉及权限系统改造。"
    )

    doc.add_heading("1. 客户线索", level=2)
    doc.add_paragraph("- 新增 / 编辑 / 删除线索（姓名、电话、公司、来源）")
    doc.add_paragraph("- 线索去重：按电话号码查重")
    doc.add_paragraph("- 批量从 Excel 导入线索")

    doc.add_heading("2. 跟进记录", level=2)
    doc.add_paragraph("- 给线索添加跟进记录（时间、内容、下次跟进时间）")
    doc.add_paragraph("- 跟进到期后发送企业微信提醒给负责销售")

    doc.add_heading("3. 商机管理", level=2)
    doc.add_paragraph("- 线索可转化为商机，商机含预估金额、阶段（接洽 / 报价 / 成交 / 失败）")
    doc.add_paragraph("- 商机阶段变更需记录审计日志")

    doc.add_heading("4. 销售看板", level=2)
    doc.add_paragraph("- 按销售、按阶段统计商机数量与金额")
    doc.add_paragraph("- 支持按时间范围筛选")

    doc.add_heading("5. 数据权限", level=2)
    doc.add_paragraph("- 销售只能看到自己名下的线索与商机")
    doc.add_paragraph("- 销售经理可查看本组全部数据")

    doc.add_heading("约束", level=2)
    doc.add_paragraph("- 技术栈：Node.js + PostgreSQL + React")
    doc.add_paragraph("- 企业微信提醒复用现有通知服务，仅需对接")

    doc.save(str(HERE / "input.docx"))
    print(f"created: {HERE / 'input.docx'}")


def make_xlsx():
    wb = Workbook()
    ws = wb.active
    ws.title = "功能清单"
    ws.append(["模块", "功能点", "说明", "优先级"])
    rows = [
        ("订单模块", "下单", "支持购物车批量下单，含库存校验", "P0"),
        ("订单模块", "订单查询", "按订单号 / 时间范围 / 状态筛选", "P0"),
        ("订单模块", "订单取消", "未发货订单可取消，库存回滚", "P0"),
        ("支付模块", "微信支付", "对接微信支付 JSAPI", "P0"),
        ("支付模块", "支付宝支付", "对接支付宝 PC 网页支付", "P0"),
        ("支付模块", "退款", "支持部分退款 + 全额退款", "P1"),
        ("物流模块", "物流跟踪", "对接快递 100 API，按订单查轨迹", "P1"),
        ("物流模块", "发货通知", "发货后短信通知买家", "P1"),
    ]
    for r in rows:
        ws.append(r)

    ws2 = wb.create_sheet("非功能要求")
    ws2.append(["要求", "指标"])
    ws2.append(["并发", "下单峰值 200 TPS"])
    ws2.append(["可用性", "99.9%"])
    ws2.append(["数据保留", "订单记录保留 3 年"])

    wb.save(str(HERE / "input.xlsx"))
    print(f"created: {HERE / 'input.xlsx'}")


if __name__ == "__main__":
    make_docx()
    make_xlsx()
