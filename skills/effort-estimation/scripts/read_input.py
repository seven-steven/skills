#!/usr/bin/env python3
"""Parse a requirements file (.docx / .xlsx / .md / .txt) into plain text on stdout.

Usage:
  python3 read_input.py <path>

Exit codes:
  0: success
  1: file not found / unsupported format
  2: missing dependency (with install hint)
"""

import sys
from pathlib import Path


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        print(
            "缺少依赖 python-docx，请安装：pip install python-docx",
            file=sys.stderr,
        )
        sys.exit(2)

    doc = Document(str(path))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def read_xlsx(path: Path) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError:
        print(
            "缺少依赖 openpyxl，请安装：pip install openpyxl",
            file=sys.stderr,
        )
        sys.exit(2)

    wb = load_workbook(filename=str(path), data_only=True, read_only=True)
    parts = []
    for sheet in wb.worksheets:
        parts.append(f"# Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = ["" if v is None else str(v).strip() for v in row]
            if any(cells):
                parts.append(" | ".join(cells))
        parts.append("")
    return "\n".join(parts)


def main():
    if len(sys.argv) != 2:
        print("Usage: python3 read_input.py <path>", file=sys.stderr)
        sys.exit(1)

    path = Path(sys.argv[1]).expanduser().resolve()
    if not path.exists():
        print(f"文件不存在: {path}", file=sys.stderr)
        sys.exit(1)

    suffix = path.suffix.lower()
    if suffix in (".md", ".txt", ""):
        sys.stdout.write(read_text(path))
    elif suffix == ".docx":
        sys.stdout.write(read_docx(path))
    elif suffix == ".xlsx":
        sys.stdout.write(read_xlsx(path))
    else:
        print(f"不支持的文件格式: {suffix}（支持 .md / .txt / .docx / .xlsx）", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
